import streamlit as st
import pandas as pd
import plotly.express as px
import requests
from datetime import datetime
import uuid
import google.generativeai as genai
from docx import Document
from io import BytesIO

# 페이지 설정
st.set_page_config(page_title="팀 예산 관리 시스템", page_icon="📊", layout="wide")

# 배포 후 발급받은 Apps Script Web App URL을 여기에 넣으세요. 
# (Streamlit Secrets를 사용해도 되지만, 간편함을 위해 바로 적어도 됩니다)
APPS_SCRIPT_URL = st.secrets.get("apps_script_url", "https://script.google.com/macros/s/AKfycbzTHUTgY1Ukz9KMuU2-gv52LhTYcyoTxVGM0-aVIOpCGWU_bmGi7yiSMCPGpgdOX-BO/exec")

@st.cache_data(ttl=5) # 5초마다 데이터 갱신
def get_data():
    try:
        response = requests.get(APPS_SCRIPT_URL)
        data = response.json()
        if not data:
            return pd.DataFrame(columns=["ID", "날짜", "팀원", "항목", "금액"])
        return pd.DataFrame(data)
    except Exception as e:
        st.error(f"데이터를 불러오는 중 오류가 발생했습니다: {e}")
        return pd.DataFrame(columns=["ID", "날짜", "팀원", "항목", "금액"])

df = get_data()

# --- 헤더 ---
st.markdown("<h1 style='text-align: center;'>📊 팀 예산 관리 시스템</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center; color: gray;'>부장님 보고용 월별 예산 취합 및 대시보드 (Apps Script 연동)</p>", unsafe_allow_html=True)
st.write("---")

# --- 탭 구성 ---
tab1, tab2, tab3 = st.tabs(["📝 데이터 입력", "📈 전체 대시보드", "📑 기간별 보고서"])

# --- TAB 1: 데이터 입력 ---
with tab1:
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.subheader("📝 내역 입력")
        with st.form("budget_form", clear_on_submit=True):
            member = st.selectbox("팀원 선택", ["부장님", "팀원1", "팀원2", "팀원3", "팀원4"])
            
            current_month = datetime.now().strftime("%Y-%m")
            month = st.text_input("해당 월 (YYYY-MM 형식)", value=current_month)
            
            category = st.selectbox("예산 항목", ["수선유지비", "비품", "개량공사"])
            amount = st.number_input("사용 금액 (원)", min_value=0, step=1000)
            
            submit_button = st.form_submit_button(label="기록 저장하기", use_container_width=True)
            
            if submit_button:
                if amount > 0:
                    new_row = [str(uuid.uuid4())[:8], month, member, category, amount]
                    payload = {
                        "action": "append",
                        "row": new_row
                    }
                    try:
                        # Apps Script로 데이터 전송 (POST)
                        response = requests.post(APPS_SCRIPT_URL, json=payload)
                        if response.status_code == 200:
                            st.success("✅ 예산 데이터가 정상적으로 기록되었습니다.")
                            st.cache_data.clear()
                            st.rerun()
                        else:
                            st.error("저장 실패 (서버 응답 오류)")
                    except Exception as e:
                        st.error(f"저장 중 오류 발생: {e}")
                else:
                    st.warning("금액은 0원보다 커야 합니다.")
                    
    with col2:
        st.subheader("📂 최근 입력 내역")
        if not df.empty:
            display_df = df.copy()
            display_df = display_df.iloc[::-1]
            # 금액에 콤마 추가 (숫자형 변환 후)
            display_df['금액'] = pd.to_numeric(display_df['금액'], errors='coerce').fillna(0)
            display_df['금액_표시'] = display_df['금액'].apply(lambda x: f"{int(x):,}원")
            
            st.dataframe(
                display_df[["날짜", "팀원", "항목", "금액_표시"]].rename(columns={"금액_표시": "금액"}), 
                use_container_width=True, 
                hide_index=True
            )
            
            if st.button("🚨 모든 데이터 초기화 (위험)", type="secondary"):
                payload = {"action": "clear"}
                requests.post(APPS_SCRIPT_URL, json=payload)
                st.cache_data.clear()
                st.rerun()
        else:
            st.info("등록된 데이터가 없습니다.")

# --- TAB 2: 전체 대시보드 ---
with tab2:
    if df.empty:
        st.warning("데이터가 없어 대시보드를 표시할 수 없습니다. 먼저 데이터를 입력해주세요.")
    else:
        df['금액'] = pd.to_numeric(df['금액'], errors='coerce').fillna(0)
        
        total_amount = df['금액'].sum()
        top_category = df.groupby('항목')['금액'].sum().idxmax()
        top_category_amount = df.groupby('항목')['금액'].sum().max()
        data_count = len(df)
        
        m1, m2, m3 = st.columns(3)
        m1.metric("전체 누적 사용액", f"{int(total_amount):,}원")
        m2.metric("최대 사용 항목", f"{top_category}", f"{int(top_category_amount):,}원 사용", delta_color="off")
        m3.metric("데이터 건수", f"{data_count}건")
        
        st.write("---")
        
        c1, c2 = st.columns(2)
        
        with c1:
            st.subheader("🏠 항목별 예산 분포")
            cat_df = df.groupby('항목', as_index=False)['금액'].sum()
            fig_pie = px.pie(cat_df, values='금액', names='항목', hole=0.4, 
                             color_discrete_sequence=['#3b82f6', '#10b981', '#8b5cf6'])
            st.plotly_chart(fig_pie, use_container_width=True)
            
        with c2:
            st.subheader("👥 팀원별 누적 사용액")
            mem_df = df.groupby('팀원', as_index=False)['금액'].sum()
            fig_bar = px.bar(mem_df, x='팀원', y='금액', text_auto='.2s', 
                             color_discrete_sequence=['#60a5fa'])
            fig_bar.update_traces(textposition='outside')
            st.plotly_chart(fig_bar, use_container_width=True)
            
        st.subheader("📅 월별/항목별 요약 테이블 (취합본)")
        pivot_df = pd.pivot_table(df, values='금액', index='날짜', columns='항목', aggfunc='sum', fill_value=0)
        pivot_df['합계'] = pivot_df.sum(axis=1)
        pivot_df = pivot_df.sort_index(ascending=False)
        
        st.dataframe(
            pivot_df.style.format("{:,.0f}"), 
            use_container_width=True
        )





# --- TAB 3: 기간별 보고서 (앱스 스크립트의 gemini 함수 연동) ---
with tab3:
    st.subheader("📑 기간별 예산 보고서 작성 (AI 어시스턴트)")
    
    if df.empty:
        st.warning("데이터가 없어 보고서를 작성할 수 없습니다. 먼저 데이터를 입력해주세요.")
    else:
        if 'reports' not in st.session_state:
            st.session_state.reports = {}

        # 1. 기간(월) 선택 기능
        unique_months = sorted(df['날짜'].unique(), reverse=True)
        selected_month = st.selectbox("보고서를 작성/조회할 기간(월)을 선택하세요:", unique_months)
        
        # 2. 선택한 기간의 예산 요약 데이터 계산
        month_df = df[df['날짜'] == selected_month]
        month_df.loc[:, '금액'] = pd.to_numeric(month_df['금액'], errors='coerce').fillna(0)
        month_total = month_df['금액'].sum()
        
        st.info(f"💡 **{selected_month}** 총 사용 금액: **{int(month_total):,}원** ({len(month_df)}건)")
        
        # 🌟 앱스 스크립트로 AI 보고서 요청하기
        if st.button("✨ AI로 부장님 보고용 초안 작성하기", type="primary"):
            if month_df.empty:
                st.warning("해당 월에는 데이터가 없습니다.")
            else:
                with st.spinner("구글 서버의 AI가 보고서를 작성 중입니다... ⏳"):
                    # 통계를 텍스트로 변환
                    cat_summary = month_df.groupby('항목')['금액'].sum()
                    cat_text = "\n".join([f"- {k}: {int(v):,}원" for k, v in cat_summary.items()])
                    
                    mem_summary = month_df.groupby('팀원')['금액'].sum()
                    mem_text = "\n".join([f"- {k}: {int(v):,}원" for k, v in mem_summary.items()])
                    
                    # 파이썬이 작성한 프롬프트 (질문 내용)
                    prompt = f"""
                    당신은 전문적이고 센스 있는 예산 관리자입니다. 
                    아래의 {selected_month} 팀 예산 사용 데이터를 바탕으로 부장님께 보고할 '월간 예산 현황 보고서' 초안을 작성해주세요.

                    [데이터 요약]
                    - 총 사용 금액: {int(month_total):,}원
                    - 총 지출 건수: {len(month_df)}건
                    
                    [항목별 지출]
                    {cat_text}
                    
                    [팀원별 지출]
                    {mem_text}

                    [작성 가이드]
                    1. 전체 요약, 주요 지출 분석, 향후 예산 운영 의견으로 구조화해주세요.
                    2. 어투는 '~습니다', '~합니다' 형태의 비즈니스 보고서 톤으로 해주세요.
                    3. 마크다운을 적절히 섞어서 가독성 좋게 작성해주세요.
                    """
                    
                    # 앱스 스크립트로 프롬프트 전송
                    payload = {
                        "action": "generate_report",
                        "prompt": prompt
                    }
                    
                    try:
                        response = requests.post(APPS_SCRIPT_URL, json=payload)
                        result = response.json()
                        
                        if result.get("status") == "success":
                            # 앱스 스크립트가 보내준 답변을 세션에 저장
                            st.session_state.reports[selected_month] = result.get("report")
                            st.rerun()
                        else:
                            st.error(f"❌ 오류가 발생했습니다: {result.get('message', '알 수 없는 오류')}")
                            
                    except Exception as e:
                        st.error(f"❌ 앱스 스크립트와 통신 실패: {e}")

        # 3. 보고서 작성 및 수정 폼
        with st.form(f"report_form_{selected_month}"):
            existing_report = st.session_state.reports.get(selected_month, "")
            
            report_content = st.text_area(
                "📝 보고서 내용 수정 및 완성", 
                value=existing_report, 
                height=400, 
                placeholder="AI가 작성한 초안을 바탕으로 내용을 수정하거나 직접 보고서를 작성하세요."
            )
            
            submit_report = st.form_submit_button("최종 보고서 저장", use_container_width=True)
            
            if submit_report:
                if report_content.strip():
                    st.session_state.reports[selected_month] = report_content
                    st.success(f"✅ {selected_month} 예산 보고서가 저장되었습니다!")
                else:
                    st.warning("보고서 내용을 입력해주세요.")
                st.rerun()

        # 4. 저장된 보고서 출력 (미리보기) 및 워드 다운로드 기능
        if st.session_state.reports.get(selected_month):
            st.write("---")
            st.subheader(f"📄 {selected_month} 예산 현황 최종 보고서")
            
            final_report_text = st.session_state.reports[selected_month]
            
            # 미리보기 화면
            st.markdown(f"""
            <div style='background-color: white; padding: 30px; border-radius: 10px; border: 1px solid #e5e7eb; box-shadow: 0 2px 4px rgba(0,0,0,0.05); margin-bottom: 20px;'>
                <pre style='font-family: inherit; white-space: pre-wrap; margin: 0;'>{final_report_text}</pre>
            </div>
            """, unsafe_allow_html=True)
            
            # --- 워드 파일(.docx) 생성 로직 ---
            doc = Document()
            # 워드 문서 제목 추가
            doc.add_heading(f"{selected_month} 예산 현황 보고서", level=0)
            # 워드 문서 본문(보고서 내용) 추가
            doc.add_paragraph(final_report_text)
            
            # 문서를 메모리 버퍼에 저장 (파일을 서버에 물리적으로 저장하지 않음)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # 다운로드 버튼 생성
            st.download_button(
                label="💾 워드 파일(.docx)로 다운로드",
                data=buffer,
                file_name=f"{selected_month}_예산보고서.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
